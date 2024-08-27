import os
import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_google_genai import ChatGoogleGenerativeAI
import google.generativeai as genai
from PyPDF2 import PdfReader
import re
from docx import Document


# Load environment variables
load_dotenv()

# Set up the API key for Google Gemini
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if GOOGLE_API_KEY is None:
    st.error("Google API key is not found. Please check your .env file.")
genai.configure(api_key=GOOGLE_API_KEY)

# Initialize Google LLM
llm = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.0)

# Streamlit App Configuration
st.set_page_config(page_title="HR Assistant", page_icon=":briefcase:", layout="centered")
st.title("AI Empowered HR Assistant 💼🤝👩🏻‍💼")


# Sidebar with navigation
with st.sidebar:
    st.header("Menu")
    choice = st.radio("Choose functionality:", ["Job Description", "HR Policy Assistance", "CV Summarize", "Extract CV Details"])


# Define function to extract text from multiple CVs
def extract_text_from_cvs(uploaded_cvs):
    texts = []
    for file in uploaded_cvs:
        file_type = file.name.split('.')[-1]
        if file_type == 'pdf':
            reader = PdfReader(file)
            text = " ".join([page.extract_text() for page in reader.pages])
        else:
            text = file.read().decode('utf-8')
        texts.append(text)
    return texts


def extract_field(extracted_text, field_name):
    pattern = rf"{field_name}: (.*)"
    match = re.search(pattern, extracted_text)
    return match.group(1) if match else ""



# Job Description Generator
if choice == "Job Description":
    st.header("Job Description Generator")

    # Inputs for JD Generator
    company_name = st.text_input("Company Name:")
    job_title = st.text_input("Job Title:")
    department = st.selectbox("Department:", ["IT", "Marketing", "HR", "Finance", "Operations"])
    job_location = st.selectbox("Job Location:", ["Mumbai", "Bangalore", "Delhi", "Hyderabad", "Remote"])
    employment_type = st.selectbox("Employment Type:", ["Full-Time", "Part-Time", "Contract"])
    experience_level = st.selectbox("Experience Level:", ["Entry-Level", "Mid-Level", "Senior-Level"])
    qualifications = st.text_area("Qualifications Required:")
    skills = st.text_area("Skills Required:")
    responsibilities = st.text_area("Key Responsibilities:")
    salary_range = st.text_input("Salary Range (in INR):")
    work_hours = st.selectbox("Work Hours:", ["9 AM - 6 PM", "Rotational", "Flexible"])

    # Generate JD Button
    if st.button("Generate Job Description"):
        # Define the prompt template for generating job descriptions
        jd_template = """
        Create a detailed job description for the following position:
        
        Company Name: {company_name}
        Job Title: {job_title}
        Department: {department}
        Location: {job_location}
        Employment Type: {employment_type}
        Experience Level: {experience_level}
        Qualifications: {qualifications}
        Skills: {skills}
        Responsibilities: {responsibilities}
        Salary Range: {salary_range}
        Work Hours: {work_hours}

        Provide the job description in a professional format.
        """

        # Prepare the prompt
        prompt = jd_template.format(
            company_name=company_name, job_title=job_title, department=department, job_location=job_location,
            employment_type=employment_type, experience_level=experience_level,
            qualifications=qualifications, skills=skills, responsibilities=responsibilities,
            salary_range=salary_range, work_hours=work_hours
        )

        # Generate job description
        try:
            chain = LLMChain(llm=llm, prompt=PromptTemplate.from_template(jd_template))
            job_description = chain.run({
                "company_name": company_name, "job_title": job_title, "department": department,
                "job_location": job_location, "employment_type": employment_type,
                "experience_level": experience_level, "qualifications": qualifications,
                "skills": skills, "responsibilities": responsibilities,
                "salary_range": salary_range, "work_hours": work_hours
            })
            st.subheader("Generated Job Description")
            st.write(job_description)
        except Exception as e:
            st.error(f"Failed to generate job description: {str(e)}")

# HR Policy Assistance
if choice == "HR Policy Assistance":
    st.header("HR Policy Assistance")

    # Inputs for HR Policy Assistance
    policy_type = st.selectbox("Select Policy Type:", ["Leave Policy", "Work From Home Policy", "Maternity Policy", "PF & Gratuity", "Harassment Policy", "Other"])
    employee_category = st.selectbox("Employee Category:", ["Permanent", "Contractual", "Interns"])

    # If the user selects "Other", display a text input field to allow them to type in the custom policy type
    if policy_type == "Other":
       custom_policy = st.text_input("Please specify the policy type:")
       policy_type = custom_policy  # Update policy_type with the custom input
    
    # Handling Leave Policy type for custom leave types
    if policy_type == "Leave Policy":
        leave_types = st.multiselect("Leave Types:", ["Sick Leave", "Earned Leave", "Maternity Leave", "Paternity Leave", "Other"])
        if "Other" in leave_types:
           custom_leave = st.text_input("Please specify the leave type:")
           leave_types.append(custom_leave)  # Add custom leave to the leave_types list
    
    # Generate HR Policy Button
    if st.button("Generate HR Policy"):
        # Define the prompt template for generating HR policies
        policy_template = """
        Generate a detailed {policy_type} for the {employee_category} category employees in an Indian context.
        Consider the following policy components:

        - {policy_type} details
        - Compliance with Indian labor laws
        - Leave entitlement (if applicable)
        - Any other relevant clauses.
        """

        # Prepare the prompt
        policy_prompt = policy_template.format(
            policy_type=policy_type,
            employee_category=employee_category
        )

        # Generate HR policy
        try:
            chain = LLMChain(llm=llm, prompt=PromptTemplate.from_template(policy_template))
            hr_policy = chain.run({"policy_type": policy_type, "employee_category": employee_category})
            st.subheader(f"Generated {policy_type}")
            st.write(hr_policy)
            
            # Provide download option for HR Policy in text format
            st.download_button(
                label="Download HR Policy as Text",
                data=hr_policy,
                file_name=f"{policy_type}_Policy.txt",
                mime="text/plain"
            )

            # Create and download Word document
            doc = Document()
            doc.add_heading(f"{policy_type} Policy", level=1)
            doc.add_paragraph(hr_policy)
            doc_file_path = f"/tmp/{policy_type}_Policy.docx"
            doc.save(doc_file_path)

            # Read the content of the Word document to offer download
            with open(doc_file_path, "rb") as file:
                st.download_button(
                    label="Download HR Policy as Word Document",
                    data=file,
                    file_name=f"{policy_type}_Policy.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Failed to generate HR policy: {str(e)}")


# CV Summarize
if choice == "CV Summarize":
    st.header("CV Summarize")

    # File upload for CV
    uploaded_cv = st.file_uploader("Upload CV (in plain text format)", type=["txt", "docx", "pdf"])

    # Generate CV Summary
    if uploaded_cv and st.button("Summarize CV"):
        try:
            # Extract text from the uploaded CV
            cv_text = extract_text_from_cvs([uploaded_cv])[0]

            # Define the prompt template for summarizing the CV
            summarize_template = """
            Summarize the following CV highlighting the key skills, experience, education, certifications, and any other relevant details:

            {cv_text}
            """

            # Generate CV summary
            chain = LLMChain(llm=llm, prompt=PromptTemplate.from_template(summarize_template))
            cv_summary = chain.run({"cv_text": cv_text})

            # Display the summarized CV
            st.subheader("CV Summary")
            st.write(cv_summary)
        except Exception as e:
            st.error(f"Failed to summarize CV: {str(e)}")

# Extract CV Details
if choice == "Extract CV Details":
    st.header("Extract CV Details")

    # File upload for CV (accepting multiple files)
    uploaded_cvs = st.file_uploader("Upload CV(s)", type=["txt", "docx", "pdf"], accept_multiple_files=True)

    # Extract CV Details Button
    if uploaded_cvs and st.button("Extract CV Details"):
        # Extract text from multiple CVs
        cvs_texts = extract_text_from_cvs(uploaded_cvs)

        # List to hold extracted details from multiple documents
        all_cv_details = []

        # Define the prompt template for extracting CV details
        extract_template = """
        You are a recruiter expert. Extract the following details from the uploaded multiple CV in a tabular format and also present in pandas dataframe:

        - Name
        - Experience (Years)
        - Education
        - Highest Qualification
        - Location
        - Current Company
        - Mobile Number
        - Email

        
        CV Text:
        {cv_text}
        """

        
        # Iterate through each CV text and extract details using LLMChain
        chain = LLMChain(llm=llm, prompt=PromptTemplate.from_template(extract_template))

        for cv_text in cvs_texts:
            extracted_details = chain.run({"cv_text": cv_text})
    
            # Debugging: Check what the LLM model outputs
            st.write("Extracted Details:", extracted_details)

            # Example extraction parsing logic based on expected output format
            # Assuming the output is a structured string, you can split it into fields
            # You'll need to parse this output based on your model's actual behavior.
            details_dict = {
               "Name": extract_field(extracted_details, "Name"),
               "Experience (Years)": extract_field(extracted_details, "Experience"),
               "Education": extract_field(extracted_details, "Education"),
               "Highest Qualification": extract_field(extracted_details, "Highest Qualification"),
               "Location": extract_field(extracted_details, "Location"),
               "Current Company": extract_field(extracted_details, "Current Company"),
               "Mobile Number": extract_field(extracted_details, "Mobile Number"),
               "Email": extract_field(extracted_details, "Email")
        }
        all_cv_details.append(details_dict)

       # st.write("Test Extracted Details:", extracted_details)


        # Convert the list of extracted details to a DataFrame
        df = pd.DataFrame(all_cv_details)

        # Display the DataFrame
        #if not df.empty:
         #  st.dataframe(df)
        #else:
       #    st.write("No details extracted.")

        # Display the table in Streamlit
        #st.subheader("Extracted CV Details from Multiple Documents")
        #st.dataframe(df)
